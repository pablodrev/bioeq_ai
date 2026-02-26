import os
import re
import logging
import requests
from typing import Optional, List, Tuple
from docx import Document
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from schemas import CriticalParametersResponse, DesignResultResponse

# --- Настройка логирования ---
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s [%(levelname)s] %(message)s',
    datefmt='%H:%M:%S'
)
logger = logging.getLogger("BioEqAssistant")

# --- Интеллектуальный модуль YandexGPT ---
class YandexGPTAssistant:
    def __init__(self, api_key: str, folder_id: str):
        self.api_key = api_key
        self.folder_id = folder_id
        self.url = "https://llm.api.cloud.yandex.net/foundationModels/v1/completion"

    def ask_expert(self, field_name: str, context_line: str, doc_context: str) -> str:
        """Интеллектуальный подбор значения через YandexGPT."""
        system_prompt = (
            "Ты — ведущий медицинский писатель, эксперт по клиническим исследованиям (CRO). "
            "Твоя специализация — биоэквивалентность (Решение ЕАЭС №85). "
            "Заполни пропуск в документе максимально профессионально и кратко. Не используй перемнные, только слова или цифры."
        )

        user_prompt = (
            f"КОНТЕКСТ ДОКУМЕНТА (начало): {doc_context[:800]}\n"
            f"СТРОКА С ПРОПУСКОМ: \"{context_line}\"\n"
            f"ЧТО НУЖНО ОПРЕДЕЛИТЬ: поле {{{field_name}}}\n\n"
            "Дай только финальный текст для вставки без кавычек и пояснений."
        )

        payload = {
            "modelUri": f"gpt://{self.folder_id}/aliceai-llm/latest",
            "completionOptions": {"stream": False, "temperature": 0.2, "maxTokens": "120"},
            "messages": [
                {"role": "system", "text": system_prompt},
                {"role": "user", "text": user_prompt}
            ]
        }
        headers = {"Authorization": f"Api-Key {self.api_key}", "x-folder-id": self.folder_id}
        
        try:
            response = requests.post(self.url, headers=headers, json=payload, timeout=15)
            response.encoding = 'utf-8'  # Явное указание UTF-8 кодирования
            response.raise_for_status()
            return response.json()['result']['alternatives'][0]['message']['text'].strip().replace('"', '')
        except Exception as e:
            logger.error(f"Ошибка LLM для {field_name}: {e}")
            return "[ТРЕБУЕТСЯ РУЧНОЙ ВВОД]"

# --- Основной процессор документа ---
class SynopsisAssistant:
    def __init__(self, drug_name_t, drug_name_r, conditions, shape, design_data: DesignResultResponse):
        self.llm = None
        self.generated_log: List[Tuple[str, str]] = [] # Для итоговой таблицы
        
        # Цветовая схема
        self.COLOR_DATA = RGBColor(0, 80, 180)    # Синий (Точные данные)
        self.COLOR_AI = RGBColor(230, 90, 0)     # Оранжевый (Сгенерировано ИИ)

        # Реестр известных переменных
        cp = design_data.critical_parameters
        self.known_vars = {
            "ТЕСТИРУЕМЫЙ ПРЕПАРАТ (T)": drug_name_t,
            "РЕФЕРЕНТНЫЙ ПРЕПАРАТ (R)": drug_name_r,
            "УСЛОВИЕ ПРИЕМА": conditions,
            "ФОРМА": shape,
            "T_MAX": f"{cp.tmax} ч" if cp.tmax else "н/д",
            "T1/2": f"{cp.t_half} ч" if cp.t_half else "н/д",
            "CV_INTRA": f"{design_data.cv_intra}%",
            "N/2": str(design_data.sample_size // 2),
            "washout_days": str(design_data.washout_days),
            "ОТМЫВОЧНЫЙ_ПЕРИОД": f"{design_data.washout_days} дней",
            "washout_days*5": str(design_data.washout_days * 5) if design_data.washout_days else "?",
            "SCREEN_FAIL_RATE": f"{design_data.screen_fail_rate}%",
            "DROPOUT_RATE": f"{design_data.dropout_rate}%"
        }

    def set_llm(self, assistant: YandexGPTAssistant):
        self.llm = assistant

    def _add_run(self, paragraph, text, color=None, bold=False, italic=False):
        run = paragraph.add_run(text)
        if color: run.font.color.rgb = color
        if bold: run.bold = True
        if italic: run.italic = True
        return run

    def fill_and_save(self, template_path: str, output_path: str):
        logger.info("Чтение шаблона и запуск анализа...")
        with open(template_path, 'r', encoding='utf-8') as f:
            full_text = f.read()

        doc = Document()
        lines = full_text.split('\n')
        
        for line in lines:
            line = line.strip()
            if not line: 
                doc.add_paragraph()
                continue

            # Обработка заголовков Markdown
            if line.startswith('## '):
                doc.add_heading(line.replace('## ', ''), level=1)
                continue
            elif line.startswith('### '):
                doc.add_heading(line.replace('### ', ''), level=2)
                continue

            p = doc.add_paragraph()
            # Разбиваем строку на части по плейсхолдерам {...}
            parts = re.split(r'(\{.*?\})', line)
            
            for part in parts:
                if part.startswith('{') and part.endswith('}'):
                    key = part.strip("{}")
                    # 1. Проверяем в словаре данных
                    if key in self.known_vars:
                        val = self.known_vars[key]
                        self._add_run(p, val, color=self.COLOR_DATA, bold=True)
                    # 2. Иначе — интеллект
                    elif self.llm:
                        logger.info(f"Интеллектуальное заполнение: {part}")
                        val = self.llm.ask_expert(key, line, full_text)
                        self._add_run(p, val, color=self.COLOR_AI, italic=True)
                        self.generated_log.append((part, val))
                    else:
                        self._add_run(p, part)
                else:
                    # Обработка жирного шрифта внутри текста **текст**
                    sub_parts = re.split(r'(\*\*.*?\*\*)', part)
                    for sp in sub_parts:
                        if sp.startswith('**') and sp.endswith('**'):
                            self._add_run(p, sp.replace('**', ''), bold=True)
                        else:
                            self._add_run(p, sp)

        self._add_review_table(doc)
        doc.save(output_path)
        logger.info(f"Файл успешно создан: {output_path}")

    def _add_review_table(self, doc):
        """Добавляет таблицу со всеми сгенерированными ИИ полями в конец документа."""
        doc.add_page_break()
        doc.add_heading("Контрольный лист сгенерированных данных", level=1)
        doc.add_paragraph("В данной таблице перечислены поля, заполненные интеллектуальным ассистентом. "
                          "Рекомендуется проверить их соответствие специфике конкретного препарата.")

        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Метка в шаблоне'
        hdr_cells[1].text = 'Сгенерированное значение'

        for placeholder, value in self.generated_log:
            row_cells = table.add_row().cells
            row_cells[0].text = placeholder
            row_cells[1].text = value
            # Красим текст в таблице в оранжевый для соответствия
            row_cells[1].paragraphs[0].runs[0].font.color.rgb = self.COLOR_AI

# --- Пример использования ---
if __name__ == "__main__":
    # 1. Загрузка данных (имитация данных из БД или калькулятора)
    crit = CriticalParametersResponse(cv_intra=26.0, tmax=2.0, t_half=14.5)
    design = DesignResultResponse(
        sample_size=28, recruitment_size=34, design_type="2-way crossover",
        cv_intra=26.0, power=0.8, alpha=0.05, dropout_rate=10.0,
        screen_fail_rate=20.0, washout_days=10.0, critical_parameters=crit
    )

    # 2. Инициализация
    api_key = os.getenv("YANDEX_GPT_API_KEY")
    folder_id = os.getenv("YANDEX_FOLDER_ID")
    
    gpt_module = YandexGPTAssistant(api_key, folder_id)
    
    synopsis = SynopsisAssistant(
        drug_name_t="Ривароксабан", 
        drug_name_r="Ксарелто®",
        conditions="натощак",
        shape="таблетки, покрытые пленочной оболочкой, 20 мг",
        design_data=design
    )
    
    synopsis.set_llm(gpt_module)
    
    # 3. Запуск
    synopsis.fill_and_save("template.md", "Готовый_Синопсис.docx")
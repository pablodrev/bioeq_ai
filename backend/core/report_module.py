"""
Report generation module with LLM-powered synopsis generation.
Generates DOCX report with intelligent document filling.
"""
import os
import re
import logging
import requests
from typing import List, Tuple
from docx import Document
from docx.shared import RGBColor

from schemas import DesignResultResponse

# --- Настройка логирования ---
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s [%(levelname)s] %(message)s',
    datefmt='%H:%M:%S'
)
logger = logging.getLogger("BioEqAssistant")

# --- Интеллектуальный модуль YandexGPT ---
class YandexGPTAssistant:
    def __init__(self, api_key: str, folder_id: str):
        self.api_key = api_key
        self.folder_id = folder_id
        self.url = "https://llm.api.cloud.yandex.net/foundationModels/v1/completion"

    def ask_expert(self, field_name: str, context_line: str, doc_context: str) -> str:
        """Интеллектуальный подбор значения через YandexGPT."""
        system_prompt = (
            "Ты — ведущий медицинский писатель, эксперт по клиническим исследованиям (CRO). "
            "Твоя специализация — биоэквивалентность (Решение ЕАЭС №85). "
            "Заполни пропуск в документе максимально профессионально и кратко. Не используй перемнные, только слова или цифры."
        )

        user_prompt = (
            f"КОНТЕКСТ ДОКУМЕНТА (начало): {doc_context[:800]}\n"
            f"СТРОКА С ПРОПУСКОМ: \"{context_line}\"\n"
            f"ЧТО НУЖНО ОПРЕДЕЛИТЬ: поле {{{field_name}}}\n\n"
            "Дай только финальный текст для вставки без кавычек и пояснений."
        )

        payload = {
            "modelUri": f"gpt://{self.folder_id}/aliceai-llm/latest",
            "completionOptions": {"stream": False, "temperature": 0.2, "maxTokens": "120"},
            "messages": [
                {"role": "system", "text": system_prompt},
                {"role": "user", "text": user_prompt}
            ]
        }
        headers = {"Authorization": f"Api-Key {self.api_key}", "x-folder-id": self.folder_id}
        
        try:
            response = requests.post(self.url, headers=headers, json=payload, timeout=15)
            response.encoding = 'utf-8'  # Явное указание UTF-8 кодирования
            response.raise_for_status()
            return response.json()['result']['alternatives'][0]['message']['text'].strip().replace('"', '')
        except Exception as e:
            logger.error(f"Ошибка LLM для {field_name}: {e}")
            return "[ТРЕБУЕТСЯ РУЧНОЙ ВВОД]"

# --- Основной процессор документа ---
class SynopsisAssistant:
    def __init__(self, drug_name_t, drug_name_r, conditions, shape, design_data: DesignResultResponse):
        self.llm = None
        self.generated_log: List[Tuple[str, str]] = [] # Для итоговой таблицы
        
        # Цветовая схема
        self.COLOR_DATA = RGBColor(0, 80, 180)    # Синий (Точные данные)
        self.COLOR_AI = RGBColor(230, 90, 0)     # Оранжевый (Сгенерировано ИИ)

        # Реестр известных переменных
        cp = design_data.critical_parameters
        self.known_vars = {
            "ТЕСТИРУЕМЫЙ ПРЕПАРАТ (T)": drug_name_t,
            "РЕФЕРЕНТНЫЙ ПРЕПАРАТ (R)": drug_name_r,
            "УСЛОВИЕ ПРИЕМА": conditions,
            "ФОРМА": shape,
            "T_MAX": f"{cp.tmax} ч" if cp.tmax else "н/д",
            "T1/2": f"{cp.t_half} ч" if cp.t_half else "н/д",
            "CV_INTRA": f"{design_data.cv_intra}%",
            "N/2": str(design_data.sample_size // 2),
            "washout_days": str(design_data.washout_days),
            "ОТМЫВОЧНЫЙ_ПЕРИОД": f"{design_data.washout_days} дней",
            "washout_days*5": str(design_data.washout_days * 5) if design_data.washout_days else "?",
            "SCREEN_FAIL_RATE": f"{design_data.screen_fail_rate}%",
            "DROPOUT_RATE": f"{design_data.dropout_rate}%"
        }

    def set_llm(self, assistant: YandexGPTAssistant):
        self.llm = assistant

    def _add_run(self, paragraph, text, color=None, bold=False, italic=False):
        run = paragraph.add_run(text)
        if color:
            run.font.color.rgb = color
        if bold:
            run.bold = True
        if italic:
            run.italic = True
        return run

    def fill_and_save(self, template_path: str, output_path: str):
        logger.info("Чтение шаблона и запуск анализа...")
        with open(template_path, 'r', encoding='utf-8') as f:
            full_text = f.read()

        doc = Document()
        lines = full_text.split('\n')
        
        for line in lines:
            line = line.strip()
            if not line: 
                doc.add_paragraph()
                continue

            # Обработка заголовков Markdown
            if line.startswith('## '):
                doc.add_heading(line.replace('## ', ''), level=1)
                continue
            elif line.startswith('### '):
                doc.add_heading(line.replace('### ', ''), level=2)
                continue

            p = doc.add_paragraph()
            # Разбиваем строку на части по плейсхолдерам {...}
            parts = re.split(r'(\{.*?\})', line)
            
            for part in parts:
                if part.startswith('{') and part.endswith('}'):
                    key = part.strip("{}")
                    # 1. Проверяем в словаре данных
                    if key in self.known_vars:
                        val = self.known_vars[key]
                        self._add_run(p, val, color=self.COLOR_DATA, bold=True)
                    # 2. Иначе — интеллект
                    elif self.llm:
                        logger.info(f"Интеллектуальное заполнение: {part}")
                        val = self.llm.ask_expert(key, line, full_text)
                        self._add_run(p, val, color=self.COLOR_AI, italic=True)
                        self.generated_log.append((part, val))
                    else:
                        self._add_run(p, part)
                else:
                    # Обработка жирного шрифта внутри текста **текст**
                    sub_parts = re.split(r'(\*\*.*?\*\*)', part)
                    for sp in sub_parts:
                        if sp.startswith('**') and sp.endswith('**'):
                            self._add_run(p, sp.replace('**', ''), bold=True)
                        else:
                            self._add_run(p, sp)

        self._add_review_table(doc)
        doc.save(output_path)
        logger.info(f"Файл успешно создан: {output_path}")

    def _add_review_table(self, doc):
        """Добавляет таблицу со всеми сгенерированными ИИ полями в конец документа."""
        doc.add_page_break()
        doc.add_heading("Контрольный лист сгенерированных данных", level=1)
        doc.add_paragraph("В данной таблице перечислены поля, заполненные интеллектуальным ассистентом. "
                          "Рекомендуется проверить их соответствие специфике конкретного препарата.")

        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Метка в шаблоне'
        hdr_cells[1].text = 'Сгенерированное значение'

        for placeholder, value in self.generated_log:
            row_cells = table.add_row().cells
            row_cells[0].text = placeholder
            row_cells[1].text = value
            # Красим текст в таблице в оранжевый для соответствия
            row_cells[1].paragraphs[0].runs[0].font.color.rgb = self.COLOR_AI


# --- Compatibility wrapper for API ---
class ReportModule:
    """Wrapper for legacy API compatibility with LLM-powered synopsis generation."""
    
    def __init__(self, db=None):
        """Initialize with optional database connection (for compatibility)."""
        self.db = db
    
    def generate_synopsis(self, project_id: str, output_path: str) -> dict:
        """
        Generate DOCX report with LLM-powered synopsis.
        Uses template.md and SynopsisAssistant for intelligent filling.
        
        Args:
            project_id: Project UUID
            output_path: Path to save DOCX file
        
        Returns:
            Dict with status and file path
        """
        try:
            if not self.db:
                return {"error": "Database connection required"}
            
            from models import DBProject
            from pathlib import Path
            
            # Fetch project
            project = self.db.query(DBProject).filter(
                DBProject.project_id == project_id
            ).first()
            
            if not project:
                return {"error": "Project not found"}
            
            # Get design parameters
            design_raw = project.design_parameters
            design = design_raw if isinstance(design_raw, dict) else {}
            
            # If design not present, attempt to generate it
            if not design:
                try:
                    from design_module import DesignModule
                    designer = DesignModule(self.db)
                    generated = designer.generate_design(project_id)
                    if isinstance(generated, dict) and not generated.get("error"):
                        design = generated
                    else:
                        logger.warning(f"Design generation: {generated.get('error') if isinstance(generated, dict) else 'unknown error'}")
                        # Return error if no design available
                        return {"error": "Cannot generate design parameters"}
                except Exception as e:
                    logger.warning(f"Failed to auto-generate design: {e}")
                    return {"error": f"Design generation failed: {str(e)}"}
            
            # Convert design dict to DesignResultResponse object if needed
            if isinstance(design, dict):
                design_obj = DesignResultResponse(**design)
            else:
                design_obj = design
            
            # Initialize SynopsisAssistant with project data
            # Extract drug names and conditions from project or use defaults
            drug_name_t = project.drug_name_t or project.inn_en or "Test Drug"
            drug_name_r = project.drug_name_r or "Reference Drug"
            conditions = getattr(project, 'administration_conditions', None) or "Fasting"
            shape = project.shape or "Not specified"
            
            synopsis = SynopsisAssistant(
                drug_name_t=drug_name_t,
                drug_name_r=drug_name_r,
                conditions=conditions,
                shape=shape,
                design_data=design_obj
            )
            
            # Setup LLM if credentials available
            api_key = os.getenv("YANDEX_GPT_API_KEY")
            folder_id = os.getenv("YANDEX_FOLDER_ID")
            
            if api_key and folder_id:
                gpt_assistant = YandexGPTAssistant(api_key, folder_id)
                synopsis.set_llm(gpt_assistant)
                logger.info(f"LLM assistant enabled for {project_id}")
            else:
                logger.info(f"LLM not configured, using static data only for {project_id}")
            
            # Determine template path
            template_dir = Path(__file__).parent
            template_path = template_dir / "template.md"
            
            if not template_path.exists():
                # Try alternative location
                template_path = Path(__file__).parent.parent / "templates" / "template.md"
            
            if not template_path.exists():
                return {"error": f"Template file not found at {template_path}"}
            
            # Generate report using SynopsisAssistant
            logger.info(f"Generating synopsis for {project_id} using template: {template_path}")
            synopsis.fill_and_save(str(template_path), output_path)
            
            logger.info(f"Report generated for {project_id} at {output_path}")
            
            return {
                "success": True,
                "file_path": output_path,
                "file_name": Path(output_path).name
            }
        
        except Exception as e:
            logger.error(f"Error generating report: {e}", exc_info=True)
            return {"error": str(e)}
